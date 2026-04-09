"""
Word Document (.docx) Formatter for ks-eye
Generates formatted Word research papers
"""


def write_docx(json_data, filepath):
    """
    Write research paper as a formatted Word document
    
    Args:
        json_data: Research JSON data dict
        filepath: Output .docx file path
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Title ---
    meta = json_data.get("metadata", {})
    title = meta.get("title", "Research Paper")
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Metadata line ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {meta.get('generated_at', 'N/A')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(json_data.get("abstract", "N/A"))

    # --- Sources Summary ---
    doc.add_heading('Sources Summary', level=1)
    sources = json_data.get("sources", {})
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Count'
    metrics = [
        ('Total Sources', sources.get('total_count', 0)),
        ('Academic Sources', sources.get('academic_sources', 0)),
        ('Web Sources', sources.get('web_sources', 0)),
        ('High Reliability', sources.get('high_reliability', 0)),
        ('Medium Reliability', sources.get('medium_reliability', 0)),
    ]
    for label, val in metrics:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(val)

    # --- Sources List ---
    source_list = sources.get("list", [])
    if source_list:
        doc.add_heading('Sources', level=2)
        for i, src in enumerate(source_list, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {src.get('title', 'Untitled')}")
            run.bold = True
            if src.get("url"):
                p.add_run(f"\n   URL: {src['url']}")
            if src.get("authors"):
                p.add_run(f"\n   Authors: {src['authors']}")
            if src.get("snippet"):
                snippet = src["snippet"][:200]
                p.add_run(f"\n   {snippet}...")
            p.add_run(f"\n   Type: {src.get('type', 'web')} | Reliability: {src.get('reliability', 'Unknown')}")
            doc.add_paragraph()  # spacer

    # --- Confidence Scores ---
    doc.add_heading('Confidence Scores', level=1)
    confidence = json_data.get("confidence_scores", {})
    conf_table = doc.add_table(rows=1, cols=3)
    conf_table.style = 'Light Shading Accent 1'
    hdr = conf_table.rows[0].cells
    hdr[0].text = 'Category'
    hdr[1].text = 'Score'
    hdr[2].text = 'Bar'
    for cat, score in confidence.items():
        row = conf_table.add_row().cells
        row[0].text = cat.replace("_", " ").title()
        row[1].text = f"{score}%"
        row[2].text = "█" * int(score / 5) + "░" * (20 - int(score / 5))

    # --- Sections ---
    sections = json_data.get("sections", [])
    for section in sections:
        doc.add_heading(section.get("title", "Section"), level=1)
        content = section.get("content", "No content available")
        doc.add_paragraph(content)

    # --- Methodology ---
    doc.add_heading('Research Methodology', level=1)
    methodology = json_data.get("methodology", {})
    doc.add_paragraph(f"Approach: {methodology.get('approach', 'N/A')}")
    doc.add_paragraph(f"Sources Used: {methodology.get('sources_used', 'N/A')}")
    doc.add_paragraph(f"Agents Involved: {methodology.get('agents_involved', 0)}")

    doc.add_paragraph("Process:", style='List Bullet')
    for step in methodology.get("process", []):
        doc.add_paragraph(step, style='List Bullet 2')

    if methodology.get("limitations"):
        doc.add_paragraph("Limitations:", style='List Bullet')
        for lim in methodology.get("limitations", []):
            doc.add_paragraph(lim, style='List Bullet 2')

    # --- Suggestions ---
    suggestions = json_data.get("suggestions", [])
    if suggestions:
        doc.add_heading('Suggested Sub-Searches', level=1)
        for i, sug in enumerate(suggestions, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {sug.get('query', 'N/A')}")
            run.bold = True
            p.add_run(f"\n   Reason: {sug.get('reason', 'N/A')}")

    # --- Limitations ---
    limitations = json_data.get("limitations", [])
    if limitations:
        doc.add_heading('Limitations', level=1)
        for lim in limitations:
            doc.add_paragraph(lim, style='List Bullet')

    # --- Footer ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Research Paper —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run = p.add_run("\nGenerated by ks-eye — KashSight Platform")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(filepath)
    return filepath
